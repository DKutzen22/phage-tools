from Bio import Entrez, SeqIO
from docx import Document
import xlsxwriter
import re
import config

#sets variables for fetching data from NCBI GenBank
Entrez.email = config.email
Entrez.tool = config.tool
Entrez.api_key = config.api_key

#retrieves genome data from NCBI GenBank
def genome_scraping(accession_num):
    #fetches the genbank data from the nucleotide database for the provided 'accession_num'
    with Entrez.efetch(db="nucleotide", 
                        id=accession_num, 
                        rettype="gb", 
                        retmode="text") as handle:
        #saves the data retrieved from genbank as a seqrecord
        seqrecord = SeqIO.read(handle, "genbank")

    return seqrecord

#organizes the reference genome into a dictionary of proteins
def protein_scraping_cds(accession_num):
    #retrieves the referenced genome
    #then creates a dict and additionally a protein counter for verification
    genome = genome_scraping(accession_num)
    sequence = genome.seq
    proteins = {}
    cds_count = 0
    #parses the seqrecord features, fetching proteins and ignoring any others
    for feature in genome.features:
        #checks that the feature is a protein coding sequence
        if feature.type != "CDS": continue
        cds_count += 1
        #variables for protein name, product, and translation
        gp = feature.qualifiers.get("note", "blank")[0].split(";")[0]
        product = feature.qualifiers.get("product", "blank")[0]
        translation = feature.qualifiers.get("translation", "blank")[0]
        #grabs the sequence of the feature by extracting it from
        #the full nucleotide sequence
        start_pos = feature.location.start
        stop_pos = feature.location.end
        gene_sequence = str(sequence[start_pos:stop_pos])
        #if the feature doesn't have a note, its named with locus_tag instead
        if not re.search(r"^gp.*", gp):
            gp = feature.qualifiers.get("locus_tag")[0]
        #places the feature in the 'proteins' dictionary with its product and translation 
        #organized by name to avoid duplication of keys
        proteins[gp] = {"Product":product, "Translation":translation, "Sequence":gene_sequence}
        if cds_count != len(proteins):
            print("Error: Not all proteins were sorted correctly")

    return proteins

#retrieves general organism data from the genome
def organism_scraping(accession_num):
    #retrieves the genome and creates a filler re.pattern for 'Description'
    genome = genome_scraping(accession_num)
    Description_pattern = re.compile(r"supercalifragilisticexpialidocious")
    #dict for various relevant facts about the organism
    description = {
        "pattern": Description_pattern,
        "Organism": {},
        "Host": {},
        "Isolation Source": {},
        "Identified by": {},
        "Location": {}
    }
    #retrieves the organism data and adds it to the 'description' dict
    for feature in genome.features:
        if feature.type == "source":
            description["Organism"]["Product"] = feature.qualifiers.get("organism", "blank")[0]
            description["Host"]["Product"] = feature.qualifiers.get("lab_host", "blank")[0]
            description["Isolation Source"]["Product"] = feature.qualifiers.get("isolation_source", "blank")[0]
            description["Identified by"]["Product"] = feature.qualifiers.get("identified_by", "blank")[0]
            description["Location"]["Product"] = feature.qualifiers.get("geo_loc_name", "blank")[0]

    return description

#organizes the proteins into a dict by class
def protein_sorting(accession_num):
    #retrieves proteins and organism data for sorting
    proteins = protein_scraping_cds(accession_num)
    organism_data = organism_scraping(accession_num)
    #re.patterns for sorting of proteins into classes
    Capsid_pattern = re.compile(r"\bhead\b|\bcapsid\b|\bcoat\b")
    Tail_pattern = re.compile(r"tail")
    Lysin_pattern = re.compile(r"lysin|lysozyme|hydrolase|amidase|endopeptidase|endoglycosidase")
    Holin_pattern = re.compile(r"holin")
    Spanin_pattern = re.compile(r"spanin")
    Toxin_pattern = re.compile(r"toxin|virulence")
    Integrase_pattern = re.compile(r"integrase")
    Adhesion_pattern = re.compile(r"immunoglobulin")
    Others_pattern = re.compile(r"superinfection|assembly|prohead|connector|portal|morphogenesis|maturation|binding|chaperone|completion|tube|measure|decor|outer|internal|inside|closure|scaffold|minor|vertex|struct|virion")
    Hypothetical_pattern = re.compile(r"\bhypothetical\sprotein\b")
    #nested dict with various classes of 
    #relevant proteins with included re.patterns
    sorted_proteins = {
        "Description": organism_data,
        "Capsid": {"pattern": Capsid_pattern},
        "Tail": {"pattern": Tail_pattern},
        "Lysin": {"pattern": Lysin_pattern},
        "Holin": {"pattern": Holin_pattern},
        "Spanin": {"pattern": Spanin_pattern},
        "Toxin": {"pattern": Toxin_pattern},
        "Integrase": {"pattern": Integrase_pattern},
        "Adhesion": {"pattern": Adhesion_pattern},
        "Others": {"pattern": Others_pattern},
        "Hypothetical": {"pattern": Hypothetical_pattern},
    }
    #iterates through proteins for sorting
    for protein in proteins:
        #retrieves the protein product and 
        #translation data stored in 'proteins'
        product = proteins.get(protein).get("Product")
        translation = proteins.get(protein).get("Translation")
        sequence = proteins.get(protein).get("Sequence")
        matched = False
        #places proteins in classes with matching re.patterns
        for class_name, sub_dict in reversed(sorted_proteins.items()):
            current_pattern = sub_dict["pattern"]
            if bool(current_pattern.search(product)):
                sorted_proteins[class_name][protein] = {"Product":product, "Translation":translation, "Sequence":sequence}
                matched = True
                #print(f"Matched: {accession_num}, {protein}, {product} to class:{class_name}")
                break
        #puts unmatched proteins into the 'Others' class
        if not matched:
            sorted_proteins["Others"][protein] = {"Product":product, "Translation":translation, "Sequence":sequence}
        
    return sorted_proteins

#makes an xlsx with pages for non-empty protein classes, 
#also includes a 'Description page
def xlsx_writer(accession_num_x, sequence_type):
    #retrieves the protein data and then creates an excel workbook
    sorted_proteins = protein_sorting(accession_num_x)
    workbook = xlsxwriter.Workbook(f"{accession_num_x} {sequence_type} - {sorted_proteins.get("Description").get("Organism").get("Product")}.xlsx")
    #iterates through the protein classes, creating worksheets for non-empty classes
    for class_name, protein_class in sorted_proteins.items():
        if len(protein_class) > 1:
            current_worksheet = workbook.add_worksheet(class_name)
            row = 0
            col = 0
            #iterates through the proteins,
            #writing their name, product, and translation.
            #skips the 'patterns' in each protein class
            for protein in protein_class:
                if protein != "pattern":
                    product = protein_class.get(protein).get("Product")
                    if sequence_type == "Proteins":
                        sequence = protein_class.get(protein).get("Translation")
                    elif sequence_type == "Genes":
                        sequence = protein_class.get(protein).get("Sequence")
                    current_worksheet.write(row, col, protein)
                    current_worksheet.write(row, col + 1, product)
                    current_worksheet.write(row, col + 2, sequence)
                    row += 1
            current_worksheet.autofit()

    workbook.close()

#creates a docx file that includes all of the protein translations from a specific class
#for a group of phages, then writes them each in fasta format with a
#naming convention of >phage_product_gp
def docx_writer(accession_num_list, protein_class):
    
    phages = accession_num_list
    target_proteins = {}
    host = organism_scraping(phages[0]).get("Host").get("Product").replace(" ", "_")

    for phage_id in phages:
        proteins = protein_sorting(phage_id)
        phage_name = proteins.get("Description").get("Organism").get("Product").rsplit("_", 1)[-1]
        target_proteins[phage_name] = {}
        for gp, protein_data in proteins.get(protein_class).items():
            if gp != "pattern":
                target_proteins[phage_name][gp] = protein_data
                    
    document = Document()

    for phage, phage_proteins in target_proteins.items():
            for gp, protein_info in phage_proteins.items():
                product = protein_info.get("Product")
                translation = protein_info.get("Translation")
                raw_name = f"{phage}_{product}_{gp}"
                name = re.sub(r'[^a-zA-Z0-9]', '_', raw_name)
                document.add_paragraph(f">{name}")
                document.add_paragraph(translation)
 
    document.save(f"{host}_phage_{protein_class}s.docx")

    
def genome_comparison_prep(phages):
    host_raw = organism_scraping(phages[0]).get("Host").get("Product")
    host = re.sub(r'[^a-zA-Z0-9]', '_', host_raw)
    with open(f"{host}_phage_genome_comparisons.fasta", "w") as file:
        for phage in phages:
            genome = genome_scraping(phage)
            organism = organism_scraping(phage).get("Organism").get("Product")
            name = re.sub(r'[^a-zA-Z0-9]', '_', organism)
            nucleotides = genome.seq
            file.write(f">{name}\n")
            file.write(f"{str(nucleotides)}\n")

    return name
